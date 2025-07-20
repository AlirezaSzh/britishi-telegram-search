from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from telethon import TelegramClient
from telethon.tl.types import PeerChannel, PeerChat
from docx import Document
from docx.shared import Inches
import os
from dotenv import load_dotenv
import asyncio
import re
from datetime import datetime
import uuid
from typing import List, Optional
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()  # Load .env file

API_ID = int(os.getenv("API_ID"))
API_HASH = os.getenv("API_HASH")
PHONE_NUMBER = os.getenv("PHONE_NUMBER")

app = FastAPI(title="Telegram Keyword Search")

# Models
class SearchRequest(BaseModel):
    telegram_link: str
    keyword: str
    limit: int = 1000  # Maximum messages to search

class SearchResult(BaseModel):
    message_id: int
    date: str
    sender: str
    content: str

# Global client instance
client = None

async def init_telegram_client():
    """Initialize Telegram client"""
    global client
    if client is None:
        client = TelegramClient('session_name', API_ID, API_HASH)
        await client.start(phone=PHONE_NUMBER)
        logger.info("Telegram client initialized")

def extract_channel_username(link: str) -> str:
    """Extract username from Telegram link"""
    patterns = [
        r't\.me/([^/?]+)',
        r'telegram\.me/([^/?]+)',
        r'@([a-zA-Z0-9_]+)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, link)
        if match:
            return match.group(1)
    
    raise ValueError("Invalid Telegram link format")

async def search_messages(channel_username: str, keyword: str, limit: int = 1000) -> List[SearchResult]:
    """Search for keyword in channel messages"""
    try:
        await init_telegram_client()
        
        # Get the channel entity
        entity = await client.get_entity(channel_username)
        
        messages = []
        keyword_lower = keyword.lower()
        
        # Iterate through messages
        async for message in client.iter_messages(entity, limit=limit):
            if message.text and keyword_lower in message.text.lower():
                sender_name = "Unknown"
                if message.sender:
                    if hasattr(message.sender, 'first_name'):
                        sender_name = message.sender.first_name or "Unknown"
                        if hasattr(message.sender, 'last_name') and message.sender.last_name:
                            sender_name += f" {message.sender.last_name}"
                    elif hasattr(message.sender, 'title'):
                        sender_name = message.sender.title
                
                messages.append(SearchResult(
                    message_id=message.id,
                    date=message.date.strftime("%Y-%m-%d %H:%M:%S") if message.date else "Unknown",
                    sender=sender_name,
                    content=message.text
                ))
        
        logger.info(f"Found {len(messages)} messages with keyword '{keyword}'")
        return messages
    
    except Exception as e:
        logger.error(f"Error searching messages: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error accessing Telegram: {str(e)}")

def create_word_document(messages: List[SearchResult], keyword: str, channel: str) -> str:
    """Create Word document with search results"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Telegram Search Results', 0)
    
    # Add search info
    doc.add_paragraph(f'Channel: {channel}')
    doc.add_paragraph(f'Keyword: "{keyword}"')
    doc.add_paragraph(f'Results found: {len(messages)}')
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('---')
    
    # Add messages
    for i, msg in enumerate(messages, 1):
        # Message header
        header = doc.add_paragraph(f'Message {i}')
        header.style = 'Heading 2'
        
        # Message details
        doc.add_paragraph(f'Date: {msg.date}')
        doc.add_paragraph(f'Sender: {msg.sender}')
        doc.add_paragraph(f'Message ID: {msg.message_id}')
        
        # Message content
        content_para = doc.add_paragraph('Content:')
        content_para.style = 'Heading 3'
        doc.add_paragraph(msg.content)
        
        # Add separator
        doc.add_paragraph('=' * 50)
    
    # Save document
    filename = f"telegram_search_{uuid.uuid4().hex[:8]}.docx"
    filepath = f"downloads/{filename}"
    
    # Create downloads directory if it doesn't exist
    os.makedirs("downloads", exist_ok=True)
    
    doc.save(filepath)
    return filename

# API Endpoints
@app.on_event("startup")
async def startup_event():
    """Initialize on startup"""
    os.makedirs("downloads", exist_ok=True)

@app.get("/", response_class=HTMLResponse)
async def read_root():
    """Serve the main HTML page"""
    html_content = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Telegram Keyword Search</title>
        <style>
            body {
                font-family: Arial, sans-serif;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background-color: #f5f5f5;
            }
            .container {
                background: white;
                padding: 30px;
                border-radius: 10px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }
            h1 {
                color: #0088cc;
                text-align: center;
                margin-bottom: 30px;
            }
            .form-group {
                margin-bottom: 20px;
            }
            label {
                display: block;
                margin-bottom: 5px;
                font-weight: bold;
                color: #333;
            }
            input[type="text"], input[type="number"] {
                width: 100%;
                padding: 12px;
                border: 1px solid #ddd;
                border-radius: 5px;
                font-size: 16px;
                box-sizing: border-box;
            }
            button {
                background-color: #0088cc;
                color: white;
                padding: 12px 30px;
                border: none;
                border-radius: 5px;
                font-size: 16px;
                cursor: pointer;
                width: 100%;
            }
            button:hover {
                background-color: #006699;
            }
            button:disabled {
                background-color: #ccc;
                cursor: not-allowed;
            }
            .loading {
                display: none;
                text-align: center;
                color: #666;
                margin-top: 20px;
            }
            .result {
                margin-top: 20px;
                padding: 15px;
                background-color: #d4edda;
                border: 1px solid #c3e6cb;
                border-radius: 5px;
                display: none;
            }
            .error {
                margin-top: 20px;
                padding: 15px;
                background-color: #f8d7da;
                border: 1px solid #f5c6cb;
                border-radius: 5px;
                display: none;
            }
            .info {
                background-color: #e7f3ff;
                border: 1px solid #b8daff;
                padding: 15px;
                border-radius: 5px;
                margin-bottom: 20px;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>🔍 Telegram Keyword Search</h1>
            
            <div class="info">
                <strong>Instructions:</strong><br>
                1. Enter a Telegram channel/group link (e.g., t.me/channelname)<br>
                2. Enter the keyword you want to search for<br>
                3. Set the message limit (max messages to search)<br>
                4. Click "Search & Download" to generate the Word document
            </div>
            
            <form id="searchForm">
                <div class="form-group">
                    <label for="telegram_link">Telegram Channel/Group Link:</label>
                    <input type="text" id="telegram_link" name="telegram_link" 
                           placeholder="https://t.me/channelname or @channelname" required>
                </div>
                
                <div class="form-group">
                    <label for="keyword">Keyword to Search:</label>
                    <input type="text" id="keyword" name="keyword" 
                           placeholder="e.g., sale, job, news" required>
                </div>
                
                <div class="form-group">
                    <label for="limit">Message Limit:</label>
                    <input type="number" id="limit" name="limit" value="1000" min="1" max="10000">
                </div>
                
                <button type="submit" id="submitBtn">🔍 Search & Download</button>
            </form>
            
            <div class="loading" id="loading">
                <p>⏳ Searching messages... This may take a few minutes.</p>
            </div>
            
            <div class="result" id="result">
                <h3>✅ Search Complete!</h3>
                <p id="resultText"></p>
                <a id="downloadLink" href="#" download>📄 Download Word Document</a>
            </div>
            
            <div class="error" id="error">
                <h3>❌ Error</h3>
                <p id="errorText"></p>
            </div>
        </div>

        <script>
            document.getElementById('searchForm').addEventListener('submit', async (e) => {
                e.preventDefault();
                
                const submitBtn = document.getElementById('submitBtn');
                const loading = document.getElementById('loading');
                const result = document.getElementById('result');
                const error = document.getElementById('error');
                
                // Reset displays
                result.style.display = 'none';
                error.style.display = 'none';
                loading.style.display = 'block';
                submitBtn.disabled = true;
                submitBtn.textContent = '🔄 Searching...';
                
                const formData = {
                    telegram_link: document.getElementById('telegram_link').value,
                    keyword: document.getElementById('keyword').value,
                    limit: parseInt(document.getElementById('limit').value)
                };
                
                try {
                    const response = await fetch('/search', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify(formData)
                    });
                    
                    if (!response.ok) {
                        const errorData = await response.json();
                        throw new Error(errorData.detail || 'Search failed');
                    }
                    
                    const data = await response.json();
                    
                    // Show result
                    document.getElementById('resultText').textContent = 
                        `Found ${data.message_count} messages containing "${formData.keyword}". Click below to download the Word document.`;
                    document.getElementById('downloadLink').href = `/download/${data.filename}`;
                    result.style.display = 'block';
                    
                } catch (err) {
                    document.getElementById('errorText').textContent = err.message;
                    error.style.display = 'block';
                } finally {
                    loading.style.display = 'none';
                    submitBtn.disabled = false;
                    submitBtn.textContent = '🔍 Search & Download';
                }
            });
        </script>
    </body>
    </html>
    """
    return HTMLResponse(content=html_content)

@app.post("/search")
async def search_telegram(request: SearchRequest):
    """Search Telegram channel/group for keyword"""
    try:
        # Extract channel username
        channel_username = extract_channel_username(request.telegram_link)
        
        # Search messages
        messages = await search_messages(channel_username, request.keyword, request.limit)
        
        if not messages:
            raise HTTPException(status_code=404, detail="No messages found with the specified keyword")
        
        # Create Word document
        filename = create_word_document(messages, request.keyword, channel_username)
        
        return {
            "success": True,
            "message_count": len(messages),
            "filename": filename,
            "channel": channel_username,
            "keyword": request.keyword
        }
    
    except Exception as e:
        logger.error(f"Search error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{filename}")
async def download_file(filename: str):
    """Download the generated Word document"""
    filepath = f"downloads/{filename}"
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)